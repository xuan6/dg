import os

from docx import Document
from docx.shared import Inches

from django.db.models import get_model
from dg.settings import MEDIA_ROOT
from dg.base_settings import STATIC_URL

from videos.scripts.exception_email import sendmail

def save_as_document(video_id):
	try:
		NonNegotiable = get_model('videos', 'NonNegotiable')
		nonnegotiables_list = NonNegotiable.objects.filter(video__id=video_id)
		Video = get_model('videos', 'Video')
		video_obj = Video.objects.get(id=video_id)

		practice_name = ''
		benefits = ''

		if video_obj.related_practice:
			if video_obj.related_practice.practice_name:
				practice_name=practice_name+str(video_obj.related_practice.practice_name)+', '
			if video_obj.related_practice.practice_sector:
				practice_name=practice_name+str(video_obj.related_practice.practice_sector)+', '
			if video_obj.related_practice.practice_subsector:
				practice_name=practice_name+str(video_obj.related_practice.practice_subsector)+', '
			if video_obj.related_practice.practice_topic:
				practice_name=practice_name+str(video_obj.related_practice.practice_topic)+', '
			if video_obj.related_practice.practice_subtopic:
				practice_name=practice_name+str(video_obj.related_practice.practice_subtopic)+', '
			if video_obj.related_practice.practice_subject:
				practice_name=practice_name+str(video_obj.related_practice.practice_subject)

		if video_obj.summary:
			benefits = benefits+str(video_obj.summary)

		logo = os.path.join("dg", "media", "output", "images", "dglogo.png")

		document = Document()
		# document.add_heading(video_obj.title, 0)
		# for obj in nonnegotiables_list:
		# 	text = obj.non_negotiable
		# 	document.add_paragraph(text)

		document.add_picture(logo, width=Inches(1.10))

		table_1 = document.add_table(rows=1, cols=1)
		table_1.style = 'TableGrid'

		hdr_cells_1 = table_1.rows[0].cells
		hdr_cells_1[0].text = 'Video Non-Negotiables Guide'
		hdr_cells_1[0].alignment = 1
		#hdr_cells_1[0].style = 'Center'

		row_cells = table_1.add_row().cells
		row_cells[0].text = 'Video Name: '+video_obj.title
		row_cells = table_1.add_row().cells
		row_cells[0].text = 'Video Category: '+practice_name
		row_cells = table_1.add_row().cells
		row_cells[0].text = 'Benefits: '+benefits
		row_cells = table_1.add_row().cells
		row_cells[0].text = 'Video:     [  ] For adoption verification    [  ] Only for understanding'

		
		table = document.add_table(rows=1, cols=2)
		table.style = 'ColorfulList-Accent3'

		hdr_cells = table.rows[0].cells
		hdr_cells[1].text = 'Physically Verifiable'
		
		count = 1
		for obj in nonnegotiables_list:
			nonnegotiable = obj.non_negotiable
			row_cells = table.add_row().cells
			row_cells[0].text = 'NON-NEGOTIABLE '+str(count)+': '+str(nonnegotiable)
			row_cells = table.add_row().cells
			row_cells[0].text = '\n'
			count = count + 1



		#table.rows[0].style = "borderColor:red; background-color:gray"


		
		#===================================================================
		# Formatting options:
		#
		# document.add_heading('Document Title', 0)
		# p = document.add_paragraph('A plain paragraph having some ')
		# p.add_run('bold').bold = True
		# p.add_run(' and some ')
		# p.add_run('italic.').italic = True
		# document.add_heading('Heading, level 1', level=1)
		# document.add_paragraph('Intense quote', style='IntenseQuote')
		# document.add_paragraph(
		# 	'first item in unordered list', style='ListBullet'
		# )
		# document.add_picture('monty-truth.png', width=Inches(1.25))
		# document.add_paragraph(
		# 	'first item in ordered list', style='ListNumber'
		# )
		# document.add_page_break()
		#===================================================================
		
		#saving file
		filename = os.path.join(MEDIA_ROOT, "non_negotiables", "%s.docx" % (video_id))
		document.save(filename)

	except Exception as e:
		error = "Video ID: "+str(video_id)+" Error: "+str(e)
		#sendmail("Exception in creating document for Non Negotiables", error)
		print error
